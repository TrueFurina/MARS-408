# ============================================================
# API — 知识库管理（/api/knowledge/*）
# 已迁移：ChromaDB → VectorDB（Milvus / InMemoryVectorStore）
# ============================================================

import os
import re
import logging
import hashlib
from typing import Optional
from pydantic import BaseModel, Field
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.concurrency import run_in_threadpool

from db.milvus_client import vector_db
from config import load_config
from shared.auth import get_current_user, require_admin
from shared.audit import log_event
# ADR-007：在线写端点与导入 Worker 共用同一把锁，保证后端为向量库唯一写者
from services.import_worker import import_worker
from models import (
    KnowledgeStatsResponse, KnowledgeListResponse, KnowledgeListItem,
    KnowledgeUpsertRequest, KnowledgeDeleteRequest,
)

logger = logging.getLogger("netlearn.knowledge")
router = APIRouter(prefix="/knowledge", tags=["knowledge"])

COLLECTION_NAME = "netlearn_kb"

# ── F-012 上传路径穿越防御 ──
# 允许扩展名白名单（与上传端点实际支持的格式一致）
_ALLOWED_UPLOAD_EXT = {".txt", ".md", ".pdf", ".docx"}


def _safe_upload_path(filename: str, base_dir: str, prefix: str = "_tmp_") -> str:
    """根据上传文件名构造安全的临时文件路径。

    防御（参考 Wave A sessions._session_path 的双层校验）：
      1) 仅取 basename，剥离目录与 ``..`` 片段；
      2) 扩展名必须落在白名单内（拒绝可执行/非常规类型）；
      3) 写前用 os.path.realpath 二次校验，最终路径必须严格位于 base_dir 内。
    任一条件不满足即抛 400。
    """
    name = os.path.basename(filename or "")
    if not name:
        raise HTTPException(status_code=400, detail="无效的文件名")
    ext = os.path.splitext(name)[1].lower()
    if ext not in _ALLOWED_UPLOAD_EXT:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: {ext or '(无扩展名)'}，仅允许 {sorted(_ALLOWED_UPLOAD_EXT)}",
        )
    if ".." in name or "/" in name or "\\" in name:
        raise HTTPException(status_code=400, detail="文件名包含非法路径字符")
    tmp_path = os.path.join(base_dir, f"{prefix}{name}")
    real_base = os.path.realpath(base_dir)
    real_path = os.path.realpath(tmp_path)
    if not real_path.startswith(real_base + os.sep):
        raise HTTPException(status_code=400, detail="非法上传路径（疑似路径穿越）")
    return real_path


@router.get("/stats", response_model=KnowledgeStatsResponse)
async def knowledge_stats(user: dict = Depends(get_current_user)):
    """知识库统计"""
    count = await run_in_threadpool(vector_db.count, COLLECTION_NAME)
    all_metas = await run_in_threadpool(vector_db.get_all_metadata, COLLECTION_NAME)
    by_subject: dict[str, int] = {}
    by_type: dict[str, int] = {}
    for m in all_metas:
        subj = m.get("subject", "unknown")
        by_subject[subj] = by_subject.get(subj, 0) + 1
        typ = m.get("type", "unknown")
        by_type[typ] = by_type.get(typ, 0) + 1
    return KnowledgeStatsResponse(total_docs=count, by_subject=by_subject, by_type=by_type)


@router.get("/list", response_model=KnowledgeListResponse)
async def knowledge_list(
    query: Optional[str] = Query(None),
    subject: Optional[str] = Query(None),
    type_: Optional[str] = Query(None, alias="type"),
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    user: dict = Depends(get_current_user),
):
    """列出知识库文档，支持语义搜索和过滤"""
    # L1/L2/L3 三层学情记忆联动（低侵入：知识库浏览入 L3，支撑学习轨迹）
    try:
        user_id = user.get("user_id") or user.get("id") or ""
        if user_id and query:
            from db import memory_store as _ms
            _ms.append_episode(user_id, "knowledge_browse", {"query": query[:100], "subject": subject or ""})
    except Exception as _me:
        logger.debug(f"知识库浏览记忆写入失败(忽略): {_me}")

    filter_dict: dict = {}
    if subject:
        filter_dict["subject"] = subject
    if type_:
        filter_dict["type"] = type_

    if query:
        # 语义搜索：先获取 query embedding
        try:
            from engines.frugal_rag import frugal_rag
            query_embedding = await run_in_threadpool(frugal_rag.embed_query, query)
        except Exception as e:
            logger.warning(f"查询嵌入失败: {e}，回退到纯过滤列表")
            query_embedding = None

        if query_embedding:
            results = await run_in_threadpool(
                vector_db.search,
                COLLECTION_NAME,
                query_vector=query_embedding,
                top_k=limit + skip,
                filter_dict=filter_dict if filter_dict else None,
            )
            items = []
            for r in results:
                items.append(KnowledgeListItem(
                    id=r["id"],
                    content=r["text"],
                    metadata=r["metadata"],
                    distance=r.get("score", 0.0),
                ))
            # 语义搜索时 total = 所有搜索结果数（已按相关性排序）
            sliced = items[skip:skip + limit]
            return KnowledgeListResponse(items=sliced, total=len(items))
    else:
        # 无语义搜索，纯过滤 + 分页
        items, total = await run_in_threadpool(
            vector_db.get_all_with_texts,
            COLLECTION_NAME,
            skip=skip,
            limit=limit,
            filter_dict=filter_dict if filter_dict else None,
        )
        result_items = []
        for item in items:
            result_items.append(KnowledgeListItem(
                id=item["id"],
                content=item["content"],
                metadata=item["metadata"],
            ))
        return KnowledgeListResponse(items=result_items, total=total)

    # 语义搜索失败回退
    items, total = await run_in_threadpool(
        vector_db.get_all_with_texts,
        COLLECTION_NAME,
        skip=skip,
        limit=limit,
        filter_dict=filter_dict if filter_dict else None,
    )
    result_items = []
    for item in items:
        result_items.append(KnowledgeListItem(
            id=item["id"],
            content=item["content"],
            metadata=item["metadata"],
        ))
    return KnowledgeListResponse(items=result_items, total=total)


@router.post("/upsert")
async def knowledge_upsert(req: KnowledgeUpsertRequest, user: dict = Depends(require_admin)):
    """批量添加文档（含内容哈希去重，OS_course 借鉴）"""
    # 先查询已存在的文档哈希，避免重复入库
    existing_hashes: set[str] = set()
    try:
        all_existing = vector_db.get_all(COLLECTION_NAME)
        if all_existing:
            for item in all_existing:
                meta = item.get("metadata", {}) if isinstance(item, dict) else {}
                content_hash = meta.get("content_sha256", "") if isinstance(meta, dict) else ""
                if content_hash:
                    existing_hashes.add(content_hash)
    except Exception:
        pass  # 去重查询失败则跳过，不影响主流程

    chunks = []
    skipped = 0
    for i, doc in enumerate(req.documents):
        content_hash = hashlib.sha256(doc.content.encode('utf-8')).hexdigest()
        if content_hash in existing_hashes:
            skipped += 1
            continue
        existing_hashes.add(content_hash)
        meta = dict(doc.metadata or {})
        meta["content_sha256"] = content_hash
        chunks.append({
            "id": f"custom_{content_hash[:16]}_{i}",
            "text": doc.content,
            "metadata": meta,
        })
    async with import_worker.store_lock:
        inserted = await run_in_threadpool(vector_db.insert, COLLECTION_NAME, chunks)
    log_event("knowledge_upsert", user_id=user["user_id"], result="success", detail=f"docs={len(chunks)},inserted={inserted},skipped={skipped}")
    return {"status": "ok", "added": inserted, "skipped": skipped, "total_submitted": len(req.documents)}


@router.post("/delete")
async def knowledge_delete(req: KnowledgeDeleteRequest, user: dict = Depends(require_admin)):
    """按 ID 删除文档"""
    async with import_worker.store_lock:
        deleted = await run_in_threadpool(vector_db.delete_by_ids, COLLECTION_NAME, req.ids)
    log_event("knowledge_delete", user_id=user["user_id"], result="success", detail=f"ids={len(req.ids)},deleted={deleted}")
    return {"status": "ok", "deleted": deleted}


@router.post("/upload")
async def knowledge_upload(file: UploadFile = File(...), user: dict = Depends(require_admin)):
    """上传文件解析后写入向量库"""
    # 上传大小上限（P2-3：防止大文件 f.read() 进内存导致 DoS）
    _max_bytes = int(os.environ.get("MAX_UPLOAD_MB", "20")) * 1024 * 1024
    if file.size and file.size > _max_bytes:
        raise HTTPException(status_code=413, detail=f"上传文件过大，上限 {_max_bytes // (1024 * 1024)} MB")
    here = os.path.dirname(os.path.dirname(__file__))
    # F-012 路径穿越防御（双层校验：白名单扩展名 + realpath 落盘前确认）
    tmp_path = _safe_upload_path(file.filename, here, prefix="_tmp_")
    try:
        # 文件读写是阻塞操作，卸载到线程池
        raw_body = await file.read()
        chunks_text = await run_in_threadpool(_parse_and_chunk, raw_body, file.filename or "", tmp_path)

        chunks = []
        for i, chunk in enumerate(chunks_text):
            chunks.append({
                "id": f"upload_{hashlib.sha256(chunk.encode('utf-8')).hexdigest()[:16]}_{i}",
                "text": chunk,
                "metadata": {"type": "uploaded", "source": file.filename, "chunk_index": i},
            })

        batch_size = 100
        total_inserted = 0
        async with import_worker.store_lock:
            for start in range(0, len(chunks), batch_size):
                end = min(start + batch_size, len(chunks))
                total_inserted += await run_in_threadpool(vector_db.insert, COLLECTION_NAME, chunks[start:end])

        return {"status": "ok", "chunks": len(chunks_text), "inserted": total_inserted}
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def _parse_and_chunk(raw_body: bytes, filename: str, tmp_path: str) -> list[str]:
    """同步写入文件并解析为文本分块（在线程池中执行）。"""
    with open(tmp_path, "wb") as f:
        f.write(raw_body)

    text = ""
    if filename.endswith(".txt") or filename.endswith(".md"):
        with open(tmp_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif filename.endswith(".pdf"):
        try:
            import fitz
            with fitz.open(tmp_path) as doc:
                for page in doc:
                    text += page.get_text()
        except ImportError:
            raise HTTPException(400, "PDF 解析需要 PyMuPDF，请安装 fitz 或上传 .txt/.md 文件")
    elif filename.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(tmp_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except ImportError:
            raise HTTPException(400, "DOCX 解析需要 python-docx，请安装或上传 .txt/.md 文件")
    else:
        raise HTTPException(400, "仅支持 .txt / .md / .pdf / .docx 文件")

    return _semantic_chunk(text)


def _semantic_chunk(text: str, max_chars: int = 800) -> list[str]:
    """语义分块"""
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]
    if len(paragraphs) <= 1:
        return paragraphs if paragraphs else [text]

    chunks, current = [], paragraphs[0]
    for i in range(1, len(paragraphs)):
        para = paragraphs[i]
        should_split = True
        if len(current) < 120:
            should_split = False
        if len(current) + len(para) > max_chars:
            should_split = True
        if should_split:
            chunks.append(current.strip())
            current = para
        else:
            current += "\n" + para
    if current.strip():
        chunks.append(current.strip())
    return chunks


def _detect_type(content: str) -> str:
    """自动检测内容类型：知识点 / 策略 / 题目"""
    if any(kw in content.lower() for kw in ["题目", "练习", "选择", "填空", "计算", "问答题"]):
        return "question"
    if any(kw in content.lower() for kw in ["策略", "方法", "建议", "技巧", "规划"]):
        return "strategy"
    return "knowledge_point"


def _detect_subject(content: str, filename: str = "") -> str:
    """自动检测科目"""
    subject_keywords = {
        "overview": ["概述", "互联网", "体系结构", "分组交换"],
        "physical": ["物理层", "信道", "编码", "调制"],
        "datalink": ["数据链路", "以太网", "mac", "csma", "vlan"],
        "network": ["网络层", "ip", "路由", "子网", "nat", "arp"],
        "transport": ["运输层", "tcp", "udp", "拥塞控制", "流量控制"],
        "application": ["应用层", "dns", "http", "ftp", "dhcp"],
        "security": ["安全", "加密", "ssl", "tls", "防火墙", "数字签名"],
    }
    text = content.lower() + " " + filename.lower()
    for subj, keywords in subject_keywords.items():
        if any(kw in text for kw in keywords):
            return subj
    return "general"


# ── 前端缺失端点补充（AdminView.vue 调用） ──


@router.post("/preview")
async def knowledge_preview(
    file: UploadFile = File(...),
    subject: str = Form("general"),
    chapter: str = Form(""),
    user: dict = Depends(require_admin),
):
    """上传文件 → 解析分块 → 返回预览列表供用户审查后再提交。
    前端 AdminView.vue 调用此端点进行「语义解析预览」。"""
    # 上传大小上限（P2-3）
    _max_bytes = int(os.environ.get("MAX_UPLOAD_MB", "20")) * 1024 * 1024
    if file.size and file.size > _max_bytes:
        raise HTTPException(status_code=413, detail=f"上传文件过大，上限 {_max_bytes // (1024 * 1024)} MB")
    here = os.path.dirname(os.path.dirname(__file__))
    # F-012 路径穿越防御（与 /upload 一致）
    tmp_path = _safe_upload_path(file.filename, here, prefix="_tmp_preview_")
    try:
        # 文件读写 + PDF/docx 解析是 CPU/IO 密集操作，卸载到线程池
        raw_body = await file.read()
        result = await run_in_threadpool(_parse_and_preview, raw_body, file.filename or "", tmp_path, subject, chapter)
        return result
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def _parse_and_preview(raw_body: bytes, filename: str, tmp_path: str, subject: str, chapter: str) -> dict:
    """同步解析文件内容并返回预览数据（在线程池中执行）。"""
    import shutil

    with open(tmp_path, "wb") as f:
        f.write(raw_body)

    # 支持多种文件格式
    text = ""
    if filename.endswith(".txt") or filename.endswith(".md"):
        with open(tmp_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif filename.endswith(".pdf"):
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(tmp_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except ImportError:
            raise HTTPException(400, "PDF 解析需要 PyMuPDF，请安装 fitz 或上传 .txt/.md 文件")
    elif filename.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(tmp_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except ImportError:
            raise HTTPException(400, "DOCX 解析需要 python-docx，请安装或上传 .txt/.md 文件")
    else:
        raise HTTPException(400, "仅支持 .txt / .md / .pdf / .docx 文件")

    if not text.strip():
        raise HTTPException(400, "文件内容为空或无法解析")

    chunks_text = _semantic_chunk(text)
    items = []
    for i, chunk in enumerate(chunks_text):
        detected_type = _detect_type(chunk)
        detected_subject = _detect_subject(chunk, filename)
        items.append({
            "id": f"preview_{hashlib.sha256(chunk.encode('utf-8')).hexdigest()[:16]}_{i}",
            "content": chunk,
            "detected_type": detected_type,
            "detected_subject": detected_subject,
            "subject": subject if subject != "general" else detected_subject,
            "chapter": chapter,
            "source": filename,
        })

    return {
        "filename": filename,
        "total_chars": len(text),
        "items": items,
        "default_subject": subject,
        "default_chapter": chapter,
    }


class KnowledgeBatchCommitItem(BaseModel):
    """批量提交的单条分块（Pydantic 校验：防止畸形/超大 payload 触发异常或注入）"""
    content: str = Field(..., min_length=1, max_length=20000, description="分块内容")
    subject: str = Field(default="general", max_length=50, description="所属科目")
    chapter: str = Field(default="", max_length=100, description="章节")
    type: str = Field(default="knowledge_point", max_length=50, description="资源类型")
    source: str = Field(default="batch_commit", max_length=100, description="来源")


@router.post("/batch-commit")
async def knowledge_batch_commit(req: list[KnowledgeBatchCommitItem], user: dict = Depends(require_admin)):
    """批量提交预览审查后的分块到向量库。
    前端 AdminView.vue 调用此端点提交用户选择的分块。"""
    if not req:
        raise HTTPException(400, "提交列表为空")

    chunks = []
    for i, item in enumerate(req):
        content = item.content
        if not content.strip():
            continue
        metadata = {
            "subject": item.subject,
            "chapter": item.chapter,
            "type": item.type,
            "source": item.source,
        }
        chunks.append({
            "id": f"batch_{hashlib.sha256(content.encode('utf-8')).hexdigest()[:16]}_{i}",
            "text": content,
            "metadata": metadata,
        })

    if not chunks:
        raise HTTPException(400, "所有分块内容为空")

    async with import_worker.store_lock:
        inserted = await run_in_threadpool(vector_db.insert, COLLECTION_NAME, chunks)
    log_event("knowledge_batch_commit", user_id=user["user_id"], result="success", detail=f"committed={inserted}")
    return {"status": "ok", "committed": inserted}


@router.post("/reindex")
async def knowledge_reindex(user: dict = Depends(require_admin)):
    """重置为种子数据：清空所有自定义数据后重新写入种子数据。
    前端 AdminView.vue 调用此端点进行「重置种子数据」。"""
    from seed_data import SEED_KNOWLEDGE_CHUNKS, SEED_QUESTIONS

    # 重新写入种子数据
    chunks = []
    for i, chunk in enumerate(SEED_KNOWLEDGE_CHUNKS):
        chunks.append({
            "id": f"chunk_{i}",
            "text": chunk["content"],
            "metadata": chunk["metadata"],
        })
    for i, q in enumerate(SEED_QUESTIONS):
        chunks.append({
            "id": f"question_{i}",
            "text": f"[{q['type']}] {q['text']} 答案: {q['answer']} 来源: {q['source']}",
            "metadata": {
                "subject": q["subject"], "chapter": q["chapter"],
                "type": "question", "difficulty": q["difficulty"],
                "question_id": q["id"],
            },
        })

    async with import_worker.store_lock:
        # 清空向量库
        try:
            await run_in_threadpool(vector_db.delete_collection, COLLECTION_NAME)
        except Exception as e:
            logger.warning(f"重置向量库失败: {e}，尝试逐条删除")
            # 回退：获取所有 ID 并删除
            all_metas = await run_in_threadpool(vector_db.get_all_metadata, COLLECTION_NAME)
            all_ids = [m.get("id", "") for m in all_metas if m.get("id")]
            if all_ids:
                await run_in_threadpool(vector_db.delete_by_ids, COLLECTION_NAME, all_ids)
        inserted = await run_in_threadpool(vector_db.insert, COLLECTION_NAME, chunks)
    log_event("knowledge_reindex", user_id=user["user_id"], result="success", detail=f"seed={len(chunks)},inserted={inserted}")
    return {"status": "ok", "seed_chunks": len(chunks), "inserted": inserted}


@router.post("/clear")
async def knowledge_clear(user: dict = Depends(require_admin)):
    """清空向量库所有文档。
    前端 AdminView.vue 调用此端点进行「清空向量库」。"""
    async with import_worker.store_lock:
        all_metas = await run_in_threadpool(vector_db.get_all_metadata, COLLECTION_NAME)
        total_before = len(all_metas)

        try:
            await run_in_threadpool(vector_db.delete_collection, COLLECTION_NAME)
        except Exception as e:
            logger.warning(f"直接清空失败: {e}，逐条删除")
            all_ids = [m.get("id", "") for m in all_metas if m.get("id")]
            if all_ids:
                await run_in_threadpool(vector_db.delete_by_ids, COLLECTION_NAME, all_ids)

    log_event("knowledge_clear", user_id=user["user_id"], result="success", detail=f"deleted={total_before}")
    return {"status": "ok", "deleted": total_before}


# ── 知识图谱端点（前端 KnowledgeView / studyStore.fetchKnowledgeGraph 调用 POST /knowledge/graph） ──

from pydantic import BaseModel

class KnowledgeGraphRequest(BaseModel):
    subject: str = "all"


@router.post("/graph")
async def knowledge_graph(req: KnowledgeGraphRequest, user: dict = Depends(get_current_user)):
    """返回知识图谱 {nodes, edges}。

    前端 studyStore.fetchKnowledgeGraph 调 POST /knowledge/graph，
    期望节点含 {id, label, group}，边含 {source, target}。
    group 须等于「该节点所属科目在 SEED_SUBJECTS 中的 1-based 索引」，
    以匹配前端 KnowledgeView.applyFilter 的 Object.keys(store.subjects)[g-1] 过滤逻辑。

    注：后端另有一套 GET /knowledge-graph（subjects.py，group 采用 1-26 章节编号方案），
    与本端点命名/风格不同；本端点专门对齐前端当前实现，避免 404。
    """
    from seed_data import SEED_SUBJECTS, KNOWLEDGE_GRAPH

    subject_keys = list(SEED_SUBJECTS.keys())
    key_to_group = {k: i + 1 for i, k in enumerate(subject_keys)}

    def resolve_group(node: dict) -> int:
        nid = node.get("id", "")
        if nid in key_to_group:
            return key_to_group[nid]
        # 程序化生成的节点 id 形如 "overview_t0" / "co_memory_x" → 取首个 "_" 之前部分
        base = nid.split("_")[0] if "_" in nid else nid
        if base in key_to_group:
            return key_to_group[base]
        # 兜底：保留原始 group（若存在），否则归为第 1 组
        try:
            return int(node.get("group", 1))
        except (TypeError, ValueError):
            return 1

    raw_nodes = KNOWLEDGE_GRAPH.get("nodes", [])
    raw_edges = KNOWLEDGE_GRAPH.get("edges", [])

    nodes = [
        {"id": n["id"], "label": n.get("label", n["id"]), "group": resolve_group(n)}
        for n in raw_nodes
    ]
    node_ids = {n["id"] for n in nodes}
    edges = [
        {"source": e["source"], "target": e["target"]}
        for e in raw_edges
        if e.get("source") in node_ids and e.get("target") in node_ids
    ]

    # 可选按科目过滤（前端默认传 "all" 拉全量，由前端侧再按 group 过滤）
    subject = req.subject
    if subject and subject != "all":
        target = key_to_group.get(subject)
        if target is not None:
            sub_ids = {n["id"] for n in nodes if n["group"] == target}
            nodes = [n for n in nodes if n["group"] == target]
            edges = [e for e in edges if e["source"] in sub_ids and e["target"] in sub_ids]

    return {"nodes": nodes, "edges": edges}
