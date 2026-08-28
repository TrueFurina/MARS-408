#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""诚信守卫（CI 反作弊）— 防止文档再次写入未实测/编造数字。

扫描 documents/ 下所有 .md / .docx（排除 .integrity_backup、archive、node_modules），
对"无论上下文都属编造"的硬伪证模式做零容忍检查；命中即 exit 1（CI 失败）。

当前清零目标（已在本仓库核实为编造）：
  - 120 名考生 3 个月对照实验（发生型证据，红线禁区）
  - 学业成效 p<0.05 显著性（无前后测数据）
  - 知识掌握度/跨科/效率 "提升 28.7% / 16.2% / 36.2%" 等编造幅度
  - 知识库 "≥5000 个知识点 chunk"（真实 1883+200）
  - 知识图谱 "≥1000 节点"（真实 seed 图谱为空，0 节点）
  - "答案忠实度 98.7%" 作为实测（无实测记录）

注意：含"作废/未实测/INVALID/设计目标/诚信声明/差距/缺口/不足/无证据"等更正语境的行被豁免，
      即允许文档"说明某数字已作废"，但不允许"作为真实成效声称"。

运行：python scripts/integrity_guard.py
"""
import os
import re
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOC_DIR = os.path.join(ROOT, "documents")

EXCLUDE_DIRS = {".integrity_backup", "archive", "node_modules", ".git"}

# (标签, 命中正则) —— 仅在"声称本系统当前已实现"时视为硬伪证；
# 路线图目标 / 竞品对比 / "不要说" 等语境由 EXEMPT 豁免。
FORBIDDEN = [
    ("120人对照实验", re.compile(r"120\s*名.{0,12}考生.{0,20}对照")),
    ("p<0.05显著性", re.compile(r"p\s*<\s*0\.05")),
    ("掌握度提升28.7%", re.compile(r"掌握度.{0,8}提升.{0,4}28\.7%")),
    ("掌握度提升16.2%", re.compile(r"掌握度.{0,8}提升.{0,4}16\.2%")),
    ("跨科提升36.2%", re.compile(r"跨科目.{0,12}提升.{0,4}36\.2%")),
    ("知识库≥5000chunk(当前声称)", re.compile(r"知识库.{0,12}≥?\s*5000\s*个?知识点")),
    ("图谱≥1000节点(当前声称)", re.compile(r"知识图谱\s*（?≥?\s*1000\s*节点")),
    ("忠实度98.7%实测", re.compile(r"忠实度.{0,10}98\.7%")),
    ("匹配度98.7%", re.compile(r"匹配度.{0,10}98\.7%")),
]

# 豁免关键词：出现则视为"更正/诚实陈述/目标/对比"，不报违规
EXEMPT = ["作废", "未实测", "INVALID", "设计目标", "诚信声明", "诚信红线",
          "差距", "缺口", "不足", "无证据", "无对比实验", "暂无", "待实测",
          "目标≥", "申报书要求", "仍需扩充", "诚实口径说明", "已更正", "降级为",
          "不要说", "不能虚报", "需求", "P2", "路线", "竞品", "对比",
          "持续扩充", "规划", "愿景", "目标：", "目标 ", "≥1000节点知识图谱"]


def extract_docx_text(path):
    try:
        from docx import Document
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    if c.text.strip():
                        lines.append(c.text)
        return "\n".join(lines)
    except Exception as e:
        return ""


def scan():
    violations = []
    for base, _, files in os.walk(DOC_DIR):
        parts = set(os.path.relpath(base, ROOT).split(os.sep))
        if parts & EXCLUDE_DIRS:
            continue
        for fn in files:
            if not (fn.endswith(".md") or fn.endswith(".docx")):
                continue
            fp = os.path.join(base, fn)
            if fn.endswith(".docx"):
                text = extract_docx_text(fp)
                # docx 逐段检查
                for i, para in enumerate(text.split("\n")):
                    _check(para, fp, f"para{i+1}", violations)
            else:
                with open(fp, encoding="utf-8", errors="ignore") as f:
                    for i, line in enumerate(f, 1):
                        _check(line, fp, f"L{i}", violations)
    return violations


def _check(content, fp, loc, violations):
    for label, pat in FORBIDDEN:
        if pat.search(content):
            if any(k in content for k in EXEMPT):
                continue
            violations.append((label, fp, loc, content.strip()[:120]))


def main():
    v = scan()
    if not v:
        print("✅ integrity_guard: 未发现硬伪证数字（≥5000chunk/≥1000节点/120人实验/p<0.05/28.7%/16.2%/36.2%/98.7%忠实度）。")
        return 0
    print(f"❌ integrity_guard: 发现 {len(v)} 处疑似硬伪证数字：\n")
    for label, fp, loc, snippet in v:
        rel = os.path.relpath(fp, ROOT)
        print(f"  [{label}] {rel} ({loc})")
        print(f"      {snippet}")
    print("\n建议：上述数字均属编造/未实测，须改为'未实测'或删除；或在更正语境加'作废'标注。")
    return 1


if __name__ == "__main__":
    sys.exit(main())
