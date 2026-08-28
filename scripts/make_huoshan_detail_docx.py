# -*- coding: utf-8 -*-
"""火山杯《作品详细介绍》docx 生成 — 数字口径与 README_火山杯-最终版.md 一致（2026-08-28 定稿）"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

NAVY = (0x12, 0x2A, 0x4A)
CYAN = (0x00, 0x9E, 0xA8)
DARK = (0x1A, 0x24, 0x2F)
GRAY = (0x5A, 0x6B, 0x7B)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.6)

def set_font(run, size=12, bold=False, color=None, name="微软雅黑"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(text, size=12, bold=False, align=None, space_after=6, color=None, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p

def h1(text):
    return para(text, size=16, bold=True, color=NAVY, space_after=8)

def h2(text):
    return para(text, size=13, bold=True, color=NAVY, space_after=6)

def bullet(text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size, False, DARK)
    return p

def table(rows, widths=None, size=11, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(v)
            set_font(r, size, header and i == 0, (0xFF, 0xFF, 0xFF) if (header and i == 0) else DARK)
            if header and i == 0:
                shd = p._p.get_or_add_pPr().makeelement(qn("w:shd"), {})
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "122A4A")
                p._p.get_or_add_pPr().append(shd)
    if widths:
        for j, w in enumerate(widths):
            for i in range(len(rows)):
                t.cell(i, j).width = Cm(w)
    return t

# ═══ 封面 ═══
para("", space_after=30)
para("2026 福建高校「火山杯」Agent 创新大赛", size=14, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
para("MARS-408", size=34, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
para("考研多智能体个性化学习系统", size=22, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
para("让 AI 从「回答问题」到「真正懂你」—— 10 节点多智能体流水线驱动的考研学习教练", size=13, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
para("赛道方向：未来学习中心（个性化学习规划 / 知识图谱 / 学习效果追踪）", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
para("覆盖 408 计算机考研四科：数据结构 · 计算机组成原理 · 操作系统 · 计算机网络", size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
doc.add_page_break()

# ═══ 一、作品概述 ═══
h1("一、作品概述")
para("MARS-408 是一个面向 408 计算机考研（数据结构、计算机组成原理、操作系统、计算机网络）的多智能体个性化学习系统。"
     "系统由 10 个各司其职的 AI Agent（协调、学情诊断、路径规划、知识检索、资源生成、出题评估、质量校验、证据核查、产物验收等）协同工作，"
     "针对每位学生的知识图谱与薄弱点，自动生成个性化学习路径与专属练习，完成「诊断 → 规划 → 讲解 → 练习 → 复盘」的完整学习闭环，"
     "让 AI 从「回答问题的工具」升级为「真正懂你的学习教练」。")
para("与普通问答助手的本质区别：系统具备任务拆解、主动规划、多轮交互与工具协同能力，"
     "能够显式暴露并裁决多 Agent 间的知识分歧，并通过证据链检索与真实大模型复核从机制上防控 AI 幻觉。", color=DARK)

# ═══ 二、痛点与场景 ═══
h1("二、痛点与场景")
para("408 计算机考研是百万级考研大军中竞争最激烈的赛道之一，备考长期面临三大困境：")
bullet("资料海量：教材、题库、网课、笔记信息过载，「该学什么」全靠个人感觉；")
bullet("无人诊断：学没学会、错在哪、下一步学什么，没有客观依据，薄弱点靠猜；")
bullet("学练割裂：刷题与讲解脱节，错题无人讲解，讲解后无人出题验证。")
para("传统 AI 助手只能「一问一答」——回答完就结束，没有诊断、没有规划、没有追踪。MARS-408 的切入点是：以多智能体协作完成学习闭环，做真正懂你的个性化学习教练。")

# ═══ 三、核心创新 ═══
h1("三、核心创新")
h2("创新 ①：多智能体共识机制，从机制上防幻觉")
para("多个 Agent 独立作答后，由基于 GoMARL 思想的加权共识引擎裁决分歧；当 Agent 间出现知识矛盾"
     "（如「三次握手 vs 四次挥手」）时，由冲突消解引擎基于知识库证据链检索，并由真实大模型复核事实后裁决。"
     "与「提示词约束防幻觉」的本质区别：分歧被显式暴露、证据被显式检索、事实被显式复核——AI 幻觉从机制上被拦截，而非依赖运气。")
h2("创新 ②：三层学习画像，越用越懂你")
bullet("短期 · 对话记忆：当前会话上下文持续跟踪，追问与纠错不丢失；")
bullet("中期 · 语义记忆：知识点掌握度动态更新，薄弱点实时感知；")
bullet("长期 · 成长画像：学习轨迹持续积累，支撑个性化路径规划与效果追踪。")
h2("创新 ③：个性化学习闭环")
para("画像驱动路径规划（薄弱点优先）→ 7 种学习资源并行生成（讲解文档、练习题、思维导图、拓展阅读、PPT 大纲、代码实操、视频脚本）"
     "→ 自动出题与批改 → 评估反馈并反哺画像。每一次提问都不是终点，而是下一次更懂你的起点。")

# ═══ 四、技术实现 ═══
h1("四、技术实现")
h2("整体架构（四层）")
table([
    ["层次", "技术", "说明"],
    ["前端层", "Vue 3 + TypeScript", "68 个页面，Vite 构建，多端多角色（学生端 / 教师看板）"],
    ["智能体层", "FastAPI + LangGraph", "10 节点多智能体流水线 · GoMARL 共识引擎 · FrugalRAG 检索引擎"],
    ["模型层", "双通道大模型自动容灾", "讯飞星火 X2（主通道）+ DeepSeek（兜底），调用失败自动切换"],
    ["数据层", "多级存储与自动降级", "Milvus / InMemory 向量库 · PostgreSQL / SQLite · Redis / 内存"],
], widths=[3.0, 4.2, 9.0])
para("")
h2("10 节点多智能体流水线（LangGraph 状态图编排）")
table([
    ["节点", "职责", "产出"],
    ["coordinator", "识别意图、全局协调、任务分派", "任务路由"],
    ["diagnostician", "学情诊断，定位薄弱知识点", "诊断报告"],
    ["planner / path_planner", "分析画像、制定分步学习计划", "学习路径"],
    ["retriever", "向量 + BM25 混合检索知识库", "证据片段"],
    ["generator_cluster", "7 种资源并行生成（含 7 个角色 Agent）", "讲解 / 习题 / 导图 / PPT / 视频脚本等"],
    ["assessor", "按知识点与难度出题、批改反馈", "练习题 / 评分"],
    ["critic / evidence_check / quality_gate", "质量审核 · 证据核查 · 产物验收", "审核报告 / 验收结论"],
], widths=[4.5, 6.2, 5.5])
para("")
h2("知识库与 FrugalRAG 自适应检索")
table([
    ["数据", "数量", "说明"],
    ["知识分片", "1883", "739 knowledge_point + 1144 knowledge_variant，覆盖四科高频考点"],
    ["练习题", "200", "选择 / 填空 / 简答，覆盖四科"],
    ["向量条目", "2083", "全部真实 E5 嵌入（768 维），0 零向量"],
    ["知识群组", "26", "按四科章节划分，供跨群冲突检测"],
], widths=[3.2, 2.0, 11.0])
para("检索链路：提问 → E5 向量检索 → BM25 全文检索 → 融合排序 → 个性化重排 → 增强生成。"
     "检索链路异常时自动降级 BM25-only（携带 _degraded 标记），绝不静默返回空结果，保证演示与真实使用场景下系统始终可用。", space_after=8)
h2("工程化水平")
bullet("前端 Vue 3 + TypeScript，68 个页面；后端 FastAPI，170+ API 端点；500+ 项自动化测试全部通过；")
bullet("讯飞星火 X2 + DeepSeek 双通道 LLM 自动容灾（超时 / 429 / 连续失败自动切换）；")
bullet("Milvus / PostgreSQL / Redis 缺失时逐级自动降级，单机即可完整运行；")
bullet("SSE 流式输出 + Agent 实时进度推送，演示效果直观流畅。")

# ═══ 五、量化实测 ═══
h1("五、量化实测（全部可复现）")
table([
    ["指标", "结果", "说明"],
    ["检索增强效果", "Recall@5 +15.8% / MRR +16.0%", "对照无重排基线，CPU 真实运行（2026-08 实测）"],
    ["检索层可回答率基线", "answerable_rate 0.533", "30 题四科验证集（eval_gold），持续回归用"],
    ["自动化测试", "500+ 项", "API 契约 / 闭环流程 / 降级路径全覆盖"],
    ["LLM 容灾", "双通道自动切换", "讯飞星火 X2 主通道 → DeepSeek 兜底"],
], widths=[3.6, 4.6, 8.0])
para("")
para("评测脚本随源码提供（py-server/experiments/），所有指标一键复现。"
     "所有成效数字均为真实运行产出，有磁盘证据与复现脚本支撑，不包含虚构的用户实验数据。", size=12, color=DARK)

# ═══ 六、落地价值 ═══
h1("六、落地价值")
bullet("学生 · 考研刚需：408 备考人群基数大、需求刚性，诊断薄弱点 → 定制路径 → 学练闭环；")
bullet("教师 · 学情管理：教师看板支持班级进度 / 掌握度 / 薄弱点聚合，从「凭经验」到「凭数据」；")
bullet("院校 · 可推广：工程完整、单机可部署、双通道容灾，可直接投入高校 / 培训机构使用。")
para("赛道定位：精准命中「未来学习中心」的三大关键词——个性化学习规划、知识图谱、学习效果追踪。")

# ═══ 七、开发工具合规声明 ═══
h1("七、开发工具合规声明")
para("本项目为真实可运行的代码工程（Vue 3 + TypeScript 前端 / FastAPI + LangGraph 后端），采用 Trae（字节跳动 AI IDE）"
     "作为开发工具完成核心模块的开发与迭代；代码仓库可在 Trae 中直接打开、构建并运行，满足 2026 福建高校「火山杯」"
     "Agent 创新大赛「基于 Trae 开发」的工具要求。关键模块：py-server/agents/graph.py（多智能体流水线）、"
     "py-server/engines/frugal_rag.py（检索引擎）、py-server/engines/gomarl.py（共识引擎）、src/（前端页面）。")

para("")
para("本材料数据口径与源码一致；量化指标均可通过随源码提供的脚本复现。", size=11, color=GRAY)

OUT = r"E:/Program/MARL/study-help-pro/submission/02_配套文档/作品详细介绍-MARS-408-最终版.docx"
doc.save(OUT)
print("saved:", OUT)
