#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""文档诚实化 v2 — docx 就地改写（仅改字，不删内容）。
按 project_审视报告-2026-08-28.md 第五节清单执行 P0-2 / P0-3。
依赖：python-docx。运行：python scripts/integrity_rewrite_docx.py
"""
import os
from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# (文件路径, [(旧串, 新串), ...])
TARGETS = [
    (os.path.join(ROOT, "documents", "基于改进GOMARL与FrugalRAG的计算机408考研个性化学习多智能体系统.docx"), [
        ("≥5000个知识点chunk", "1883 个知识点 chunk（739 知识点 + 1144 变式）"),
        ("知识图谱（≥1000个节点）", "知识图谱（当前 seed 图谱为空，规模随用户数据积累）"),
        ("知识图谱（≥1000节点，≥3000边）", "知识图谱（当前 seed 图谱为空，先修关系未落库）"),
    ]),
    (os.path.join(ROOT, "附件1： 闽江学院大学生创新创业训练计划项目中期检查报告书.docx"), [
        ("知识点先修图谱（613 节点）", "知识点先修图谱（当前 seed 图谱为空）"),
        ("节点数 613（与代码知识图谱实际节点数一致）",
         "节点数 0（代码 KNOWLEDGE_GRAPH 实际为 0 节点；文档此前‘613 节点与代码一致’表述不实，已更正）"),
    ]),
    (os.path.join(ROOT, "附件1：福建省大学生创新创业训练计划结题验收报告书（草稿）.docx"), [
        ("知识图谱（613 节点 / 609 条先修关系）",
         "知识图谱（当前 seed 图谱为空，先修关系未落库；文档此前‘613 节点/609 边’表述不实，已更正）"),
    ]),
]


def replace_in_runs(para, old, new):
    """在段落的 run 级别做字符串替换；若跨 run 则退化为段落级重建。"""
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # 跨 run 兜底：整段文本包含但分散在多个 run
    full = "".join(r.text for r in para.runs)
    if old in full:
        # 重建：保留首 run 文本，清空其余
        idx = 0
        rem = full.replace(old, new)
        for run in para.runs:
            L = len(run.text)
            run.text = rem[idx:idx + L]
            idx += L
        return True
    return False


def main():
    total = 0
    for path, subs in TARGETS:
        if not os.path.exists(path):
            print("SKIP (missing):", path)
            continue
        doc = Document(path)
        for para in doc.paragraphs:
            for old, new in subs:
                if replace_in_runs(para, old, new):
                    total += 1
                    print(f"  REPLACED [{os.path.basename(path)}]: {old!r} -> {new!r}")
        # 表格内文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in subs:
                            if replace_in_runs(para, old, new):
                                total += 1
                                print(f"  REPLACED[table] [{os.path.basename(path)}]: {old!r}")
        doc.save(path)
        print("SAVED:", path)
    print(f"\nTOTAL replacements: {total}")


if __name__ == "__main__":
    main()
